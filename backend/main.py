"""
NyayaAI — FastAPI Backend Server
AI-Powered Legal Document Generator for India
"""
import os, json, time, hashlib, uuid
from datetime import datetime, timedelta
from typing import Optional, List, Dict, Any
from contextlib import asynccontextmanager

from fastapi import FastAPI, HTTPException, Depends, status, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel
try:
    from pydantic import EmailStr
except ImportError:
    EmailStr = str  # fallback if email-validator not installed
import anthropic

# Local imports
from rag_engine import rag_engine
from clause_db import CLAUSE_DATABASE, ALL_CLAUSES, INDIAN_ACTS

# ─── Startup / Shutdown ────────────────────────────────────────────────────────
@asynccontextmanager
async def lifespan(app: FastAPI):
    print("[NyayaAI] Starting up…")
    rag_engine.initialize()
    print("[NyayaAI] RAG engine ready")
    yield
    print("[NyayaAI] Shutting down…")

app = FastAPI(
    title="NyayaAI API",
    description="AI-Powered Legal Document Generator for India",
    version="1.0.0",
    lifespan=lifespan
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── Anthropic Client ──────────────────────────────────────────────────────────
def get_claude():
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        raise HTTPException(503, "ANTHROPIC_API_KEY not configured")
    return anthropic.Anthropic(api_key=api_key)

# ─── In-Memory Store (replace with DB in production) ──────────────────────────
USERS: Dict[str, dict] = {
    "demo@nyayaai.in": {
        "id": "user_001", "name": "Arjun Rao", "email": "demo@nyayaai.in",
        "password_hash": hashlib.sha256("demo123".encode()).hexdigest(),
        "plan": "pro", "profession": "Business Owner",
        "city": "Bengaluru", "state": "Karnataka",
        "created_at": "2026-01-15", "queries_used": 0, "queries_limit": 1000,
    }
}
DOCUMENTS: Dict[str, dict] = {}
SESSIONS: Dict[str, str] = {}  # token -> user_id

# ─── Pydantic Models ───────────────────────────────────────────────────────────
class LoginRequest(BaseModel):
    email: str
    password: str

class SignupRequest(BaseModel):
    name: str
    email: str
    password: str
    profession: str = "Individual"

class GenerateDocRequest(BaseModel):
    doc_type: str
    fields: Dict[str, str]
    state: str = "Karnataka"
    language: str = "english"
    use_rag: bool = True

class ChatRequest(BaseModel):
    message: str
    history: List[Dict[str, str]] = []

class ClauseExplainRequest(BaseModel):
    clause_text: str
    clause_title: str = ""

class ClauseSearchRequest(BaseModel):
    query: str
    doc_type: str = ""
    top_k: int = 5

class RiskAnalysisRequest(BaseModel):
    doc_type: str
    document_text: str

class SaveDocRequest(BaseModel):
    doc_type: str
    title: str
    content: str
    fields: Dict[str, str] = {}
    state: str = "Karnataka"
    status: str = "draft"

# ─── Auth Helper ──────────────────────────────────────────────────────────────
def get_current_user(request: Request) -> Optional[dict]:
    token = request.headers.get("Authorization", "").replace("Bearer ", "")
    user_id = SESSIONS.get(token)
    if not user_id:
        return None
    for u in USERS.values():
        if u["id"] == user_id:
            return u
    return None

def require_auth(request: Request) -> dict:
    user = get_current_user(request)
    if not user:
        raise HTTPException(status.HTTP_401_UNAUTHORIZED, "Authentication required")
    return user

# ─── Document Templates ────────────────────────────────────────────────────────
DOC_TEMPLATES = {
    "nda": {
        "name": "Non-Disclosure Agreement",
        "subtitle": "Confidentiality Agreement | Indian Contract Act, 1872",
        "fields": [
            {"id": "party1", "label": "Disclosing Party", "placeholder": "Full name or company name", "required": True},
            {"id": "party2", "label": "Receiving Party", "placeholder": "Full name or company name", "required": True},
            {"id": "purpose", "label": "Purpose of Disclosure", "placeholder": "E.g. Evaluating a business partnership", "required": True},
            {"id": "duration", "label": "Agreement Duration", "type": "select", "options": ["1 Year", "2 Years", "3 Years", "5 Years", "Perpetual"]},
            {"id": "penalty", "label": "Breach Penalty (₹)", "placeholder": "E.g. 5,00,000"},
            {"id": "party1_address", "label": "Disclosing Party Address", "placeholder": "Full address with city and pincode"},
            {"id": "party2_address", "label": "Receiving Party Address", "placeholder": "Full address with city and pincode"},
        ],
        "recommendations": ["Non-Compete Clause", "Non-Solicitation Clause", "Return of Information Clause", "Injunctive Relief Clause"],
        "icon": "📄"
    },
    "rental": {
        "name": "Rental / Leave & License Agreement",
        "subtitle": "Residential Lease | Transfer of Property Act, 1882",
        "fields": [
            {"id": "landlord", "label": "Landlord / Licensor Name", "placeholder": "Full legal name", "required": True},
            {"id": "tenant", "label": "Tenant / Licensee Name", "placeholder": "Full legal name", "required": True},
            {"id": "property_address", "label": "Property Address", "placeholder": "Full address with pincode", "required": True},
            {"id": "rent", "label": "Monthly Rent (₹)", "placeholder": "E.g. 25,000", "required": True},
            {"id": "deposit", "label": "Security Deposit (₹)", "placeholder": "E.g. 1,00,000", "required": True},
            {"id": "lockin", "label": "Lock-in Period", "type": "select", "options": ["3 Months", "6 Months", "11 Months", "1 Year"]},
            {"id": "notice", "label": "Notice Period", "type": "select", "options": ["15 Days", "30 Days", "60 Days", "90 Days"]},
            {"id": "start_date", "label": "Commencement Date", "placeholder": "DD/MM/YYYY"},
        ],
        "recommendations": ["Maintenance Responsibility Clause", "Utility Bills Clause", "Pet Policy", "Subletting Restriction"],
        "icon": "🏠"
    },
    "employment": {
        "name": "Employment Contract",
        "subtitle": "Full-Time Employment Agreement | Indian Labour Laws",
        "fields": [
            {"id": "employer", "label": "Employer (Company Name)", "placeholder": "Full legal company name", "required": True},
            {"id": "employee", "label": "Employee Name", "placeholder": "Full legal name", "required": True},
            {"id": "designation", "label": "Designation / Role", "placeholder": "E.g. Senior Software Engineer", "required": True},
            {"id": "ctc", "label": "Annual CTC (₹)", "placeholder": "E.g. 18,00,000", "required": True},
            {"id": "department", "label": "Department", "placeholder": "E.g. Engineering"},
            {"id": "joining_date", "label": "Date of Joining", "placeholder": "DD/MM/YYYY"},
            {"id": "probation", "label": "Probation Period", "type": "select", "options": ["No Probation", "1 Month", "3 Months", "6 Months"]},
            {"id": "notice", "label": "Notice Period", "type": "select", "options": ["15 Days", "1 Month", "2 Months", "3 Months"]},
            {"id": "work_location", "label": "Work Location", "placeholder": "City, State"},
        ],
        "recommendations": ["IP Assignment Clause", "Confidentiality Clause", "Non-Compete Clause", "ESOP Clause", "Moonlighting Policy"],
        "icon": "💼"
    },
    "internship": {
        "name": "Internship Agreement",
        "subtitle": "Internship / Apprenticeship Agreement | Indian Contract Act",
        "fields": [
            {"id": "company", "label": "Company Name", "required": True},
            {"id": "intern", "label": "Intern Name", "required": True},
            {"id": "role", "label": "Internship Role", "placeholder": "E.g. Product Design Intern"},
            {"id": "duration", "label": "Duration", "type": "select", "options": ["1 Month", "2 Months", "3 Months", "6 Months"]},
            {"id": "stipend", "label": "Monthly Stipend (₹)", "placeholder": "E.g. 15,000 or Unpaid"},
            {"id": "start_date", "label": "Start Date", "placeholder": "DD/MM/YYYY"},
        ],
        "recommendations": ["IP Assignment", "Confidentiality Clause", "Work From Home Policy"],
        "icon": "🎓"
    },
    "freelance": {
        "name": "Freelance / Independent Contractor Agreement",
        "subtitle": "Project-based Service Contract | Indian Contract Act, 1872",
        "fields": [
            {"id": "client", "label": "Client Name / Company", "required": True},
            {"id": "freelancer", "label": "Freelancer / Contractor Name", "required": True},
            {"id": "project", "label": "Project / Scope of Work", "placeholder": "Brief description"},
            {"id": "fee", "label": "Total Project Fee (₹)", "required": True},
            {"id": "deadline", "label": "Project Deadline", "placeholder": "DD/MM/YYYY"},
            {"id": "payment_terms", "label": "Payment Terms", "type": "select", "options": ["50% advance + 50% delivery", "Monthly", "Milestone-based", "Full on delivery"]},
            {"id": "revisions", "label": "Number of Revisions Included", "placeholder": "E.g. 2"},
        ],
        "recommendations": ["Kill Fee Clause", "IP Ownership Clause", "Non-Solicitation Clause", "GST Clause"],
        "icon": "🤝"
    },
    "service": {
        "name": "Service Agreement",
        "subtitle": "B2B Service Contract | Indian Contract Act, 1872",
        "fields": [
            {"id": "service_provider", "label": "Service Provider", "required": True},
            {"id": "client", "label": "Client", "required": True},
            {"id": "services", "label": "Description of Services", "required": True},
            {"id": "fee", "label": "Service Fee (₹)", "required": True},
            {"id": "payment_terms", "label": "Payment Terms", "type": "select", "options": ["Monthly", "Quarterly", "Milestone-based", "Annual"]},
            {"id": "duration", "label": "Contract Duration", "type": "select", "options": ["6 Months", "1 Year", "2 Years", "Ongoing"]},
            {"id": "sla", "label": "SLA / Response Time", "placeholder": "E.g. 24 hours"},
        ],
        "recommendations": ["Limitation of Liability", "Indemnification Clause", "SLA Penalty Clause"],
        "icon": "⚙️"
    },
    "partnership": {
        "name": "Partnership Deed",
        "subtitle": "LLP/Firm Partnership Deed | Indian Partnership Act, 1932",
        "fields": [
            {"id": "firm_name", "label": "Firm / LLP Name", "required": True},
            {"id": "partner1", "label": "Partner 1 Name", "required": True},
            {"id": "partner2", "label": "Partner 2 Name", "required": True},
            {"id": "business_nature", "label": "Nature of Business", "required": True},
            {"id": "capital1", "label": "Capital Contribution - Partner 1 (₹)"},
            {"id": "capital2", "label": "Capital Contribution - Partner 2 (₹)"},
            {"id": "profit_ratio", "label": "Profit Sharing Ratio", "placeholder": "E.g. 50:50 or 60:40"},
            {"id": "address", "label": "Principal Place of Business"},
        ],
        "recommendations": ["Dispute Resolution Clause", "Retirement/Admission Clause", "Bank Account Operation Clause"],
        "icon": "🏢"
    },
    "mou": {
        "name": "Memorandum of Understanding (MoU)",
        "subtitle": "Non-Binding MoU | Indian Contract Act, 1872",
        "fields": [
            {"id": "party1", "label": "Party 1", "required": True},
            {"id": "party2", "label": "Party 2", "required": True},
            {"id": "purpose", "label": "Purpose of MoU", "required": True},
            {"id": "duration", "label": "Duration", "type": "select", "options": ["6 Months", "1 Year", "2 Years"]},
            {"id": "obligations1", "label": "Obligations of Party 1"},
            {"id": "obligations2", "label": "Obligations of Party 2"},
        ],
        "recommendations": ["Binding Clauses List", "Termination Clause", "Confidentiality Clause"],
        "icon": "📋"
    },
    "legalnotice": {
        "name": "Legal Notice",
        "subtitle": "Formal Legal Notice | Indian Contract Act, 1872",
        "fields": [
            {"id": "sender", "label": "Sender / Your Name (or Advocate)", "required": True},
            {"id": "recipient", "label": "Recipient Name / Company", "required": True},
            {"id": "subject", "label": "Subject / Nature of Dispute", "required": True},
            {"id": "amount", "label": "Amount in Dispute (₹)", "placeholder": "If applicable"},
            {"id": "deadline", "label": "Response Deadline", "type": "select", "options": ["7 Days", "15 Days", "30 Days", "60 Days"]},
            {"id": "facts", "label": "Brief Facts", "placeholder": "Describe the issue in 2-3 sentences"},
        ],
        "recommendations": ["Without Prejudice Note", "Legal Cost Recovery Clause", "Arbitration Clause"],
        "icon": "🔔"
    },
    "vendor": {
        "name": "Vendor Agreement",
        "subtitle": "Goods/Services Vendor Contract | Indian Contract Act, 1872",
        "fields": [
            {"id": "buyer", "label": "Buyer / Company Name", "required": True},
            {"id": "vendor", "label": "Vendor Name", "required": True},
            {"id": "goods_services", "label": "Goods / Services", "required": True},
            {"id": "value", "label": "Contract Value (₹)"},
            {"id": "payment_terms", "label": "Payment Terms", "type": "select", "options": ["Net 30", "Net 60", "Advance Payment", "COD", "Milestone-based"]},
            {"id": "delivery", "label": "Delivery Terms / Timeline"},
        ],
        "recommendations": ["Warranty Clause", "Defect Liability Period", "Penalty Clause"],
        "icon": "📦"
    },
    "consultancy": {
        "name": "Consultancy Agreement",
        "subtitle": "Independent Consultant Contract | Indian Contract Act, 1872",
        "fields": [
            {"id": "company", "label": "Company / Client Name", "required": True},
            {"id": "consultant", "label": "Consultant Name", "required": True},
            {"id": "scope", "label": "Scope of Consultancy", "required": True},
            {"id": "fee", "label": "Consultancy Fee (₹)", "required": True},
            {"id": "fee_type", "label": "Fee Type", "type": "select", "options": ["Monthly Retainer", "Per Day Rate", "Project-based", "Hourly Rate"]},
            {"id": "duration", "label": "Duration", "type": "select", "options": ["3 Months", "6 Months", "1 Year", "Ongoing"]},
        ],
        "recommendations": ["IP Assignment Clause", "Confidentiality Clause", "Non-Compete Clause"],
        "icon": "💡"
    },
    "poa": {
        "name": "Power of Attorney",
        "subtitle": "General / Specific PoA | Powers of Attorney Act, 1882",
        "fields": [
            {"id": "principal", "label": "Principal (Grantor) Name", "required": True},
            {"id": "attorney", "label": "Attorney (Agent) Name", "required": True},
            {"id": "poa_type", "label": "PoA Type", "type": "select", "options": ["General PoA", "Special / Specific PoA", "Durable PoA"]},
            {"id": "powers", "label": "Specific Powers Granted", "placeholder": "E.g. property sale, banking, legal matters"},
            {"id": "duration", "label": "Duration", "type": "select", "options": ["1 Year", "2 Years", "5 Years", "Until Revoked"]},
        ],
        "recommendations": ["Revocation Rights Clause", "Accountability Clause"],
        "icon": "⚖️"
    },
    "affidavit": {
        "name": "Affidavit",
        "subtitle": "Sworn Affidavit | Indian Evidence Act, 1872",
        "fields": [
            {"id": "deponent", "label": "Deponent (Your Full Name)", "required": True},
            {"id": "father_name", "label": "Father's Name", "required": True},
            {"id": "address", "label": "Deponent's Address", "required": True},
            {"id": "purpose", "label": "Purpose of Affidavit", "required": True, "placeholder": "E.g. Name change, Address proof, Court submission"},
            {"id": "statements", "label": "Statements (Key Facts)", "placeholder": "List the facts you are affirming"},
        ],
        "recommendations": ["Notarization Note", "Penalty for False Declaration"],
        "icon": "📜"
    },
    "privacy": {
        "name": "Privacy Policy",
        "subtitle": "Website/App Privacy Policy | IT Act, 2000 | PDPB Compliant",
        "fields": [
            {"id": "company", "label": "Company / App Name", "required": True},
            {"id": "website", "label": "Website URL"},
            {"id": "data_types", "label": "Types of Data Collected", "placeholder": "E.g. Name, email, payment info, device data"},
            {"id": "data_use", "label": "How Data is Used", "placeholder": "E.g. Service delivery, analytics, marketing"},
            {"id": "contact_email", "label": "Privacy Contact Email"},
        ],
        "recommendations": ["Data Retention Policy", "Third-Party Sharing Clause", "Cookies Policy"],
        "icon": "🔒"
    },
    "tnc": {
        "name": "Terms & Conditions",
        "subtitle": "Website/App T&C | IT Act, 2000 | Consumer Protection Act",
        "fields": [
            {"id": "company", "label": "Company / App Name", "required": True},
            {"id": "website", "label": "Website / App URL"},
            {"id": "services", "label": "Description of Services", "required": True},
            {"id": "payment_policy", "label": "Payment & Refund Policy"},
            {"id": "contact_email", "label": "Support Email"},
        ],
        "recommendations": ["Refund Policy Clause", "Limitation of Liability", "GDPR Alignment Note"],
        "icon": "📑"
    },
    "gst": {
        "name": "GST Tax Invoice",
        "subtitle": "GST-Compliant Invoice | CGST Act, 2017",
        "fields": [
            {"id": "supplier", "label": "Supplier / Company Name", "required": True},
            {"id": "gstin_supplier", "label": "Supplier GSTIN", "placeholder": "15-digit GSTIN"},
            {"id": "buyer", "label": "Buyer / Customer Name", "required": True},
            {"id": "gstin_buyer", "label": "Buyer GSTIN (if registered)"},
            {"id": "invoice_no", "label": "Invoice Number", "required": True},
            {"id": "items", "label": "Items / Services", "placeholder": "Item name, HSN code, quantity, rate"},
            {"id": "gst_type", "label": "GST Type", "type": "select", "options": ["CGST + SGST (Intra-state)", "IGST (Inter-state)"]},
        ],
        "recommendations": ["E-Invoice Note", "Place of Supply Clause", "TDS Note"],
        "icon": "💰"
    },
}

# ─── Routes ────────────────────────────────────────────────────────────────────

@app.get("/")
def root():
    return {"app": "NyayaAI", "version": "1.0.0", "status": "operational",
            "rag": rag_engine._initialized}

@app.get("/api/health")
def health():
    return {"status": "ok", "rag_initialized": rag_engine._initialized,
            "clauses_indexed": len(ALL_CLAUSES), "timestamp": datetime.utcnow().isoformat()}

# ── Auth ──────────────────────────────────────────────────────────────────────
@app.post("/api/auth/login")
def login(req: LoginRequest):
    user = USERS.get(req.email)
    if not user:
        raise HTTPException(400, "Invalid email or password")
    expected = hashlib.sha256(req.password.encode()).hexdigest()
    if user["password_hash"] != expected:
        raise HTTPException(400, "Invalid email or password")
    token = str(uuid.uuid4())
    SESSIONS[token] = user["id"]
    return {"token": token, "user": {k: v for k, v in user.items() if k != "password_hash"}}

@app.post("/api/auth/signup")
def signup(req: SignupRequest):
    if req.email in USERS:
        raise HTTPException(400, "Email already registered")
    user_id = f"user_{str(uuid.uuid4())[:8]}"
    USERS[req.email] = {
        "id": user_id, "name": req.name, "email": req.email,
        "password_hash": hashlib.sha256(req.password.encode()).hexdigest(),
        "plan": "free", "profession": req.profession,
        "city": "", "state": "Karnataka",
        "created_at": datetime.utcnow().strftime("%Y-%m-%d"),
        "queries_used": 0, "queries_limit": 50,
    }
    token = str(uuid.uuid4())
    SESSIONS[token] = user_id
    user = {k: v for k, v in USERS[req.email].items() if k != "password_hash"}
    return {"token": token, "user": user}

@app.post("/api/auth/google")
def google_login():
    """Simulate Google OAuth for demo purposes."""
    email = "arjun.rao@gmail.com"
    if email not in USERS:
        USERS[email] = {
            "id": f"user_google_{str(uuid.uuid4())[:8]}",
            "name": "Arjun Rao", "email": email,
            "password_hash": "", "plan": "pro", "profession": "Business Owner",
            "city": "Bengaluru", "state": "Karnataka",
            "created_at": datetime.utcnow().strftime("%Y-%m-%d"),
            "queries_used": 0, "queries_limit": 1000,
        }
    token = str(uuid.uuid4())
    SESSIONS[token] = USERS[email]["id"]
    return {"token": token, "user": {k: v for k, v in USERS[email].items() if k != "password_hash"}}

@app.get("/api/auth/me")
def get_me(request: Request):
    user = get_current_user(request)
    if not user:
        raise HTTPException(401, "Not authenticated")
    return {k: v for k, v in user.items() if k != "password_hash"}

# ── Templates ─────────────────────────────────────────────────────────────────
@app.get("/api/templates")
def get_templates():
    return [{"type": k, **{fk: fv for fk, fv in v.items() if fk != "fields"}}
            for k, v in DOC_TEMPLATES.items()]

@app.get("/api/templates/{doc_type}")
def get_template(doc_type: str):
    t = DOC_TEMPLATES.get(doc_type)
    if not t:
        raise HTTPException(404, f"Template '{doc_type}' not found")
    return {"type": doc_type, **t}

# ── RAG Clause Retrieval ──────────────────────────────────────────────────────
@app.post("/api/rag/search")
def rag_search(req: ClauseSearchRequest):
    results = rag_engine.retrieve(req.query, req.doc_type or None, req.top_k)
    return {"query": req.query, "results": results, "count": len(results),
            "method": "faiss" if rag_engine._initialized else "keyword"}

@app.get("/api/rag/clauses/{doc_type}")
def get_doc_clauses(doc_type: str):
    clauses = rag_engine.get_clauses_for_document(doc_type)
    return {"doc_type": doc_type, "clauses": clauses, "count": len(clauses)}

# ── Document Generation ───────────────────────────────────────────────────────
@app.post("/api/generate")
async def generate_document(req: GenerateDocRequest, request: Request):
    template = DOC_TEMPLATES.get(req.doc_type)
    if not template:
        raise HTTPException(400, f"Unknown document type: {req.doc_type}")

    try:
        claude = get_claude()
    except HTTPException:
        return _fallback_generate(req, template)

    # Build RAG context
    rag_context = ""
    if req.use_rag:
        rag_context = rag_engine.build_rag_context(req.doc_type, req.fields)

    fields_str = "\n".join(f"- {k}: {v}" for k, v in req.fields.items() if v)
    lang_instruction = "Generate the complete document in Hindi (Devanagari script)." if req.language == "hindi" else "Generate in English."

    system_prompt = """You are an expert Indian legal document drafter with 20+ years of experience. 
You have deep knowledge of the Indian Contract Act 1872, Transfer of Property Act 1882, 
Companies Act 2013, Labour Laws, LLP Act 2008, IT Act 2000, and all major Indian statutes.
Your documents are court-ready, professionally formatted, and fully compliant with Indian law."""

    user_prompt = f"""Draft a complete, professional {template['name']} for India.

{rag_context}

DOCUMENT DETAILS PROVIDED BY USER:
{fields_str}

State/Jurisdiction: {req.state}
{lang_instruction}

INSTRUCTIONS:
1. Use ALL the user-provided details (names, amounts, dates) verbatim in the document
2. Include all standard clauses for {template['name']} under Indian law
3. Format with clear numbered sections (1., 1.1, etc.)
4. Include recitals/whereas clauses at the beginning  
5. Add complete signature blocks with date lines at the end
6. Reference specific Indian Acts/Sections where applicable
7. Make it court-ready and professionally formatted
8. Include the governing law as {req.state}, India
9. Add a disclaimer: "This document was generated by NyayaAI for informational purposes. Please have it reviewed by a qualified legal professional."

Generate the COMPLETE document now, starting directly with the title:"""

    try:
        message = claude.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2500,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}]
        )
        doc_text = message.content[0].text

        # Update user query count
        user = get_current_user(request)
        if user:
            user["queries_used"] = user.get("queries_used", 0) + 1

        # Parse clauses from generated text for editor
        clauses = _extract_clauses(doc_text)
        risks = _analyze_risks(req.doc_type, doc_text)

        return {
            "success": True,
            "document": doc_text,
            "doc_type": req.doc_type,
            "template_name": template["name"],
            "clauses": clauses,
            "risks": risks,
            "rag_used": req.use_rag,
            "rag_clauses_retrieved": len(rag_engine.get_clauses_for_document(req.doc_type)),
            "generated_at": datetime.utcnow().isoformat(),
        }
    except Exception as e:
        return _fallback_generate(req, template)


def _fallback_generate(req: GenerateDocRequest, template: dict) -> dict:
    """Fallback document generation using clause database."""
    clauses = rag_engine.get_clauses_for_document(req.doc_type)[:6]
    f = req.fields
    party1 = f.get("party1") or f.get("party1_name") or f.get("landlord") or f.get("employer") or f.get("client") or f.get("company") or "[Party One]"
    party2 = f.get("party2") or f.get("party2_name") or f.get("tenant") or f.get("employee") or f.get("freelancer") or f.get("vendor") or "[Party Two]"
    today = datetime.utcnow().strftime("%d %B %Y")
    state = req.state

    doc = f"""{template['name'].upper()}

THIS {template['name'].upper()} ("Agreement") is entered into as of {f.get('start_date') or f.get('joining_date') or today} ("Effective Date") by and between:

{party1} ("Party One" / "First Party"), AND
{party2} ("Party Two" / "Second Party").

WHEREAS, the parties desire to set forth the terms and conditions of their agreement as follows:

"""
    for clause in clauses:
        doc += f"{clause['title'].upper()}\n\n{clause['text']}\n\n"

    doc += f"""GOVERNING LAW AND JURISDICTION

This Agreement shall be governed by and construed in accordance with the laws of India. Any dispute arising out of or in connection with this Agreement shall be subject to the exclusive jurisdiction of courts in {state}, India.

ENTIRE AGREEMENT

This Agreement constitutes the entire agreement between the parties with respect to the subject matter hereof.

IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.

PARTY ONE: {party1}
Signature: _________________________
Name: _________________________
Date: _________________________
Place: _________________________

PARTY TWO: {party2}
Signature: _________________________
Name: _________________________
Date: _________________________
Place: _________________________

WITNESS 1: _________________________    WITNESS 2: _________________________

---
Disclaimer: This document was generated by NyayaAI for informational purposes. Please have it reviewed by a qualified legal professional before execution."""

    return {
        "success": True, "document": doc, "doc_type": req.doc_type,
        "template_name": template["name"],
        "clauses": [{"title": c["title"], "text": c["text"], "risk": c.get("risk_level","low")} for c in clauses],
        "risks": _analyze_risks(req.doc_type, doc),
        "rag_used": True, "rag_clauses_retrieved": len(clauses),
        "generated_at": datetime.utcnow().isoformat(), "fallback": True
    }


def _extract_clauses(text: str) -> List[Dict]:
    """Extract numbered clauses from generated document text."""
    import re
    clauses = []
    lines = text.split('\n')
    current_title = None
    current_text = []
    for line in lines:
        line = line.strip()
        if not line:
            if current_title and current_text:
                clauses.append({"title": current_title, "text": " ".join(current_text), "risk": "low"})
                current_title = None
                current_text = []
            continue
        # Match "1. TITLE", "1.1 Title", "SECTION 1:", etc.
        heading = re.match(r'^(\d+\.[\d.]*|SECTION\s+\d+[:.:]?)\s+(.+)$', line, re.IGNORECASE)
        if heading and len(line) < 100:
            if current_title and current_text:
                clauses.append({"title": current_title, "text": " ".join(current_text), "risk": "low"})
            current_title = line
            current_text = []
        elif current_title:
            if not line.startswith("Signature") and not line.startswith("Date"):
                current_text.append(line)
    if current_title and current_text:
        clauses.append({"title": current_title, "text": " ".join(current_text), "risk": "low"})
    # Assign risk levels
    high_risk_keywords = ["penalty", "breach", "termination", "liability", "indemnif", "ip assign", "intellectual property"]
    medium_risk_keywords = ["probation", "notice", "security deposit", "lock-in", "non-compete", "confidential"]
    for c in clauses:
        text_lower = (c["title"] + c["text"]).lower()
        if any(k in text_lower for k in high_risk_keywords):
            c["risk"] = "high"
        elif any(k in text_lower for k in medium_risk_keywords):
            c["risk"] = "medium"
    return clauses[:12]


def _analyze_risks(doc_type: str, doc_text: str) -> List[Dict]:
    """Simple rule-based risk analysis."""
    risks = []
    text_lower = doc_text.lower()
    risk_rules = [
        ("arbitration" not in text_lower, "high", "Missing Dispute Resolution Clause", "No arbitration or mediation clause found. Consider adding an ADR clause (Arbitration and Conciliation Act, 1996) to avoid costly litigation."),
        ("force majeure" not in text_lower, "medium", "No Force Majeure Clause", "The document lacks a force majeure provision. This could expose parties to liability for events beyond their control (Section 56, Indian Contract Act)."),
        ("governing law" not in text_lower and "jurisdiction" not in text_lower, "high", "Governing Law Not Specified", "The jurisdiction and governing law are not clearly defined. This is essential for enforceability under Indian law."),
        (doc_type == "nda" and "injunctive" not in text_lower, "medium", "No Injunctive Relief Clause", "NDA without injunctive relief provision may limit enforcement options in case of breach."),
        (doc_type == "employment" and "non-compete" in text_lower, "high", "Non-Compete Enforceability Risk", "Post-employment non-compete clauses are generally unenforceable under Section 27, Indian Contract Act, 1872. Consider limiting scope."),
        (doc_type == "rental" and "stamp" not in text_lower, "high", "Stamp Duty Not Mentioned", "Rental agreement must be on appropriate stamp paper and registered. Check state-specific stamp duty requirements."),
        (doc_type == "freelance" and "gst" not in text_lower.replace("cgst",""), "medium", "GST Not Addressed", "If the freelancer's annual turnover exceeds ₹20L, GST registration and charging is mandatory under CGST Act, 2017."),
        ("entire agreement" not in text_lower and "whole agreement" not in text_lower, "low", "No Entire Agreement Clause", "Consider adding an entire agreement (merger) clause to prevent disputes about prior oral representations."),
    ]
    for condition, level, title, desc in risk_rules:
        if condition:
            risks.append({"level": level, "title": title, "desc": desc})
    if not risks:
        risks.append({"level": "low", "title": "Document Looks Good", "desc": "No major legal risks detected. Standard clauses are present. Always have a qualified advocate review before execution."})
    return risks

# ── AI Chat ───────────────────────────────────────────────────────────────────
@app.post("/api/chat")
async def chat(req: ChatRequest, request: Request):
    try:
        claude = get_claude()
    except HTTPException:
        return {"reply": _fallback_chat(req.message)}

    system_prompt = """You are NyayaAI, an expert AI legal assistant specialised in Indian law. You have comprehensive knowledge of:
- Indian Contract Act, 1872
- Transfer of Property Act, 1882 (rental, property laws)
- IT Act, 2000 (digital contracts, privacy, cybercrime)
- Companies Act, 2013 (corporate law, directors, governance)
- LLP Act, 2008 (limited liability partnerships)
- Industrial Disputes Act, 1947 (employment, labour rights)
- Payment of Wages Act, 1936; EPF Act, 1952; ESI Act, 1948
- Consumer Protection Act, 2019
- CGST Act, 2017 (GST and taxation)
- Specific Relief Act, 1963
- Arbitration and Conciliation Act, 1996
- Indian Partnership Act, 1932
- Patents Act, 1970; Copyright Act, 1957
- Code of Civil Procedure, 1908
- Negotiable Instruments Act, 1881 (cheque bounce)

Response style:
1. Be helpful, accurate, and concise (2-4 paragraphs max)
2. Reference specific Indian Acts and Sections when relevant
3. Use **bold** for key legal terms
4. When explaining clauses, mention: what it means, why it matters, risk if removed
5. End with: "⚠️ Note: For your specific situation, please consult a qualified advocate."
6. Use bullet points for lists
7. Be friendly but professional — you're a trusted legal guide"""

    messages = []
    for h in req.history[-6:]:
        if h.get("role") in ["user", "assistant"] and h.get("content"):
            messages.append({"role": h["role"], "content": h["content"]})
    messages.append({"role": "user", "content": req.message})

    try:
        response = claude.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=800,
            system=system_prompt,
            messages=messages
        )
        reply = response.content[0].text
        user = get_current_user(request)
        if user:
            user["queries_used"] = user.get("queries_used", 0) + 1
        return {"reply": reply, "success": True}
    except Exception as e:
        return {"reply": _fallback_chat(req.message), "success": False}


def _fallback_chat(q: str) -> str:
    q_lower = q.lower()
    if "indemnity" in q_lower:
        return "**Indemnity** (Indian Contract Act 1872, Sections 124-125) means one party agrees to compensate the other for losses arising from specified events.\n\n**Why it matters:** It allocates risk clearly between parties and provides a contractual remedy without needing to prove negligence separately.\n\n**Risk if removed:** You'd have to prove all elements of a tort or breach claim independently, which is more time-consuming and uncertain.\n\n⚠️ Note: For your specific situation, please consult a qualified advocate."
    if "force majeure" in q_lower:
        return "**Force Majeure** is governed by **Section 56 of the Indian Contract Act, 1872** (Doctrine of Frustration). It excuses a party from performance when extraordinary events beyond their control make performance impossible.\n\n**Examples:** Natural disasters, pandemics (upheld by Indian courts post-COVID), government orders, war.\n\n**Key requirement:** The event must be truly unforeseeable and make performance *impossible* — not just more expensive.\n\n⚠️ Note: For your specific situation, please consult a qualified advocate."
    if "non-compete" in q_lower:
        return "**Non-compete clauses** in India have a complex legal status:\n\n- **During employment:** Generally enforceable\n- **Post-employment:** Usually **NOT enforceable** under **Section 27, Indian Contract Act, 1872**, which voids agreements in restraint of trade\n- **Exceptions:** The Supreme Court has upheld limited non-competes where they protect genuine business interests with reasonable scope and duration\n\n**Practical advice:** If you must include it, keep the geographic scope and duration narrow, and consider making it reasonable.\n\n⚠️ Note: For your specific situation, please consult a qualified advocate."
    return "That's an important legal question under Indian law. The answer depends on the specific facts and applicable statutes. Generally, Indian courts follow the principle of reasonableness and public policy when interpreting such clauses.\n\nI'd recommend reviewing the relevant provisions and consulting with a qualified advocate for your specific situation.\n\n⚠️ Note: For your specific situation, please consult a qualified advocate."

# ── Clause Explanation ────────────────────────────────────────────────────────
@app.post("/api/clause/explain")
async def explain_clause(req: ClauseExplainRequest):
    try:
        claude = get_claude()
        response = claude.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=400,
            messages=[{"role": "user", "content": f"""Explain this Indian legal clause in simple plain English. Be concise (3-4 sentences).
Format:
**Plain English:** [simple explanation]
**Why it matters:** [importance]
**Risk if removed:** [consequence]
**Relevant Act:** [Indian law reference]

Clause Title: {req.clause_title}
Clause: {req.clause_text}"""}]
        )
        return {"explanation": response.content[0].text, "success": True}
    except:
        return {"explanation": f"**Plain English:** This clause establishes legal obligations between parties regarding {req.clause_title.lower() or 'the specified matter'}.\n\n**Why it matters:** It provides clarity on rights and responsibilities, reducing potential disputes.\n\n**Risk if removed:** Without this clause, the parties may have undefined obligations, leading to potential legal disputes.\n\n**Relevant Act:** Indian Contract Act, 1872", "success": False}

# ── Risk Analysis ─────────────────────────────────────────────────────────────
@app.post("/api/risk-analysis")
async def risk_analysis(req: RiskAnalysisRequest, request: Request):
    risks = _analyze_risks(req.doc_type, req.document_text)
    try:
        claude = get_claude()
        response = claude.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=600,
            messages=[{"role": "user", "content": f"""Analyze this {req.doc_type} document for legal risks under Indian law. 
Identify 3-5 specific risks or missing clauses. Return as JSON array:
[{{"level":"high/medium/low","title":"Risk Title","desc":"Description"}}]
Return ONLY the JSON array.

Document excerpt: {req.document_text[:1500]}"""}]
        )
        raw = response.content[0].text.strip()
        clean = raw.replace("```json", "").replace("```", "").strip()
        ai_risks = json.loads(clean)
        return {"risks": ai_risks, "success": True, "ai_powered": True}
    except:
        return {"risks": risks, "success": True, "ai_powered": False}

# ── Documents CRUD ─────────────────────────────────────────────────────────────
@app.post("/api/documents")
def save_document(req: SaveDocRequest, request: Request):
    user = get_current_user(request)
    doc_id = str(uuid.uuid4())
    DOCUMENTS[doc_id] = {
        "id": doc_id, "user_id": user["id"] if user else "guest",
        "doc_type": req.doc_type, "title": req.title, "content": req.content,
        "fields": req.fields, "state": req.state, "status": req.status,
        "version": 1, "created_at": datetime.utcnow().isoformat(),
        "updated_at": datetime.utcnow().isoformat(), "pages": max(1, len(req.content) // 2000),
    }
    return {"success": True, "doc_id": doc_id, "document": DOCUMENTS[doc_id]}

@app.get("/api/documents")
def list_documents(request: Request):
    user = get_current_user(request)
    user_id = user["id"] if user else "guest"
    docs = [d for d in DOCUMENTS.values() if d.get("user_id") == user_id]
    docs.sort(key=lambda x: x["created_at"], reverse=True)
    return {"documents": docs, "total": len(docs)}

@app.get("/api/documents/{doc_id}")
def get_document(doc_id: str):
    doc = DOCUMENTS.get(doc_id)
    if not doc:
        raise HTTPException(404, "Document not found")
    return doc

@app.put("/api/documents/{doc_id}")
def update_document(doc_id: str, req: SaveDocRequest):
    doc = DOCUMENTS.get(doc_id)
    if not doc:
        raise HTTPException(404, "Document not found")
    doc.update({"content": req.content, "status": req.status,
                "version": doc["version"] + 1, "updated_at": datetime.utcnow().isoformat()})
    return {"success": True, "document": doc}

@app.delete("/api/documents/{doc_id}")
def delete_document(doc_id: str):
    if doc_id not in DOCUMENTS:
        raise HTTPException(404, "Document not found")
    del DOCUMENTS[doc_id]
    return {"success": True}

# ── Export helpers ─────────────────────────────────────────────────────────────
def _sanitize(text: str) -> str:
    """Normalise common Unicode punctuation to ASCII-safe equivalents."""
    replacements = {
        "\u2013": "-", "\u2014": "--", "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"', "\u2026": "...", "\u20b9": "Rs.",
        "\u00a0": " ", "\u2022": "-", "\u2192": "->",
        "\u00ae": "(R)", "\u00a9": "(C)", "\u00bd": "1/2", "\u00bc": "1/4",
    }
    for uni, asc in replacements.items():
        text = text.replace(uni, asc)
    return text

# ── Export ─────────────────────────────────────────────────────────────────────
@app.get("/api/export/pdf/{doc_id}")
def export_pdf(doc_id: str):
    doc = DOCUMENTS.get(doc_id)
    if not doc:
        raise HTTPException(404, "Document not found")
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
        import io

        buf = io.BytesIO()
        doc_pdf = SimpleDocTemplate(
            buf, pagesize=A4,
            leftMargin=20*mm, rightMargin=20*mm,
            topMargin=22*mm, bottomMargin=22*mm,
        )

        styles = getSampleStyleSheet()
        # Custom styles
        title_style = ParagraphStyle("NTitle", parent=styles["Title"],
            fontSize=15, leading=20, spaceAfter=4, alignment=TA_CENTER,
            textColor=colors.HexColor("#0d0f14"), fontName="Helvetica-Bold")
        meta_style = ParagraphStyle("NMeta", parent=styles["Normal"],
            fontSize=8, leading=11, spaceAfter=8, alignment=TA_CENTER,
            textColor=colors.HexColor("#888888"))
        heading_style = ParagraphStyle("NHead", parent=styles["Normal"],
            fontSize=11, leading=14, spaceBefore=8, spaceAfter=3,
            fontName="Helvetica-Bold", textColor=colors.HexColor("#0d0f14"))
        subhead_style = ParagraphStyle("NSub", parent=styles["Normal"],
            fontSize=10, leading=13, spaceBefore=4, spaceAfter=2,
            fontName="Helvetica-Bold", textColor=colors.HexColor("#222222"))
        body_style = ParagraphStyle("NBody", parent=styles["Normal"],
            fontSize=10, leading=15, spaceAfter=3, alignment=TA_JUSTIFY,
            textColor=colors.HexColor("#1a1a1a"))
        footer_style = ParagraphStyle("NFooter", parent=styles["Normal"],
            fontSize=7.5, leading=10, alignment=TA_CENTER,
            textColor=colors.HexColor("#aaaaaa"), fontName="Helvetica-Oblique")

        story = []

        # Title
        story.append(Paragraph(_sanitize(doc.get("title", "Legal Document")), title_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#c9a84c"), spaceAfter=6))
        story.append(Paragraph(
            f"Generated by NyayaAI &nbsp;|&nbsp; {doc.get('created_at','')[:10]} &nbsp;|&nbsp; {doc.get('state','')} Jurisdiction",
            meta_style))
        story.append(Spacer(1, 6*mm))

        # Body lines
        for line in doc["content"].split("\n"):
            line = _sanitize(line.strip())
            if not line:
                story.append(Spacer(1, 2*mm))
                continue
            # Escape HTML special chars for reportlab
            safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            # ALL CAPS section headings
            if line.isupper() and 3 < len(line) < 90 and not line.startswith("("):
                story.append(Spacer(1, 2*mm))
                story.append(Paragraph(safe, heading_style))
            # Numbered headings
            elif len(line) < 90 and line.split(".")[0].isdigit():
                story.append(Paragraph(safe, subhead_style))
            else:
                story.append(Paragraph(safe, body_style))

        # Footer disclaimer
        story.append(Spacer(1, 8*mm))
        story.append(HRFlowable(width="100%", thickness=0.4, color=colors.HexColor("#cccccc"), spaceAfter=4))
        story.append(Paragraph(
            "This document was generated by NyayaAI for informational purposes. "
            "Please review with a qualified advocate before execution.",
            footer_style))

        doc_pdf.build(story)
        buf.seek(0)

        safe_title = "".join(c for c in doc.get("title", "Document") if c.isalnum() or c in " _-").strip().replace(" ", "_")
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.pdf"'}
        )
    except Exception as e:
        raise HTTPException(500, f"PDF generation failed: {e}")
    doc = DOCUMENTS.get(doc_id)
    if not doc:
        raise HTTPException(404, "Document not found")
    try:
        from fpdf import FPDF
        import io

        pdf = FPDF(orientation="P", unit="mm", format="A4")
        pdf.set_margins(15, 15, 15)
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # ── Header ──────────────────────────────────────────
        pdf.set_font("Helvetica", "B", 15)
        pdf.set_text_color(13, 15, 20)
        title = _sanitize(doc.get("title", "Legal Document"))
        pdf.multi_cell(0, 9, title, align="C")
        pdf.ln(1)

        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(120, 120, 120)
        meta = f"Generated by NyayaAI  |  {doc.get('created_at','')[:10]}  |  {doc.get('state','')} Jurisdiction"
        pdf.cell(0, 6, _sanitize(meta), ln=True, align="C")

        # Divider line
        pdf.ln(2)
        pdf.set_draw_color(200, 168, 76)
        pdf.set_line_width(0.5)
        pdf.line(20, pdf.get_y(), 190, pdf.get_y())
        pdf.ln(5)

        def safe_cell(text, height=6, bold=False, size=10):
            """Render text safely, chunking any extreme long tokens."""
            pdf.set_font("Helvetica", "B" if bold else "", size)
            # Break any single token longer than 80 chars (e.g. long URLs/codes)
            words = []
            for w in text.split():
                while len(w) > 80:
                    words.append(w[:80])
                    w = w[80:]
                words.append(w)
            safe_text = " ".join(words)
            try:
                pdf.multi_cell(0, height, safe_text)
            except Exception:
                # Last resort: write char by char truncated
                pdf.multi_cell(0, height, safe_text[:200] + "...")

        # ── Body ─────────────────────────────────────────────
        pdf.set_text_color(30, 30, 30)
        for line in doc["content"].split("\n"):
            line = _sanitize(line.strip())
            if not line:
                pdf.ln(2)
                continue
            # Section headings: ALL CAPS lines under 80 chars
            if line.isupper() and 3 < len(line) < 80 and not line.startswith("("):
                pdf.ln(3)
                pdf.set_text_color(13, 15, 20)
                safe_cell(line, height=7, bold=True, size=11)
                pdf.set_text_color(30, 30, 30)
                pdf.ln(1)
            # Numbered headings like "1." "1.1"
            elif len(line) < 80 and line[:3].replace(".", "").strip().isdigit():
                pdf.set_text_color(13, 15, 20)
                safe_cell(line, height=6, bold=True, size=10)
                pdf.set_text_color(30, 30, 30)
            else:
                safe_cell(line, height=6, bold=False, size=10)

        # ── Footer on every page ─────────────────────────────
        pdf.set_y(-15)
        pdf.set_font("Helvetica", "I", 7)
        pdf.set_text_color(160, 160, 160)
        pdf.cell(0, 5,
            "This document was generated by NyayaAI. Please review with a qualified advocate before execution.",
            align="C")

        raw = pdf.output(dest="S")
        # fpdf2 >= 2.x returns bytearray; older versions return str
        if isinstance(raw, (bytes, bytearray)):
            pdf_bytes = bytes(raw)
        else:
            pdf_bytes = raw.encode("latin-1")
        safe_title = "".join(c for c in doc.get("title","Document") if c.isalnum() or c in " _-").strip().replace(" ", "_")
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.pdf"'}
        )
    except Exception as e:
        raise HTTPException(500, f"PDF generation failed: {e}")

@app.get("/api/export/docx/{doc_id}")
def export_docx(doc_id: str):
    doc = DOCUMENTS.get(doc_id)
    if not doc:
        raise HTTPException(404, "Document not found")
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        import io

        d = DocxDocument()

        # Page margins
        for section in d.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        # Title
        title_para = d.add_heading(doc.get("title", "Legal Document"), level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(13, 15, 20)

        # Metadata line
        meta = d.add_paragraph(
            f"Generated by NyayaAI  |  {doc.get('created_at','')[:10]}  |  {doc.get('state','')} Jurisdiction"
        )
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in meta.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(120, 120, 120)
            run.font.italic = True

        d.add_paragraph()  # spacer

        # Body
        for line in doc["content"].split("\n"):
            line = line.strip()
            if not line:
                d.add_paragraph()
                continue
            # Section headings
            if line.isupper() and 3 < len(line) < 80 and not line.startswith("("):
                h = d.add_heading(line, level=2)
                for run in h.runs:
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(13, 15, 20)
            # Numbered headings
            elif len(line) < 80 and line[:3].replace(".", "").strip().isdigit():
                p = d.add_paragraph(line)
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            else:
                p = d.add_paragraph(line)
                for run in p.runs:
                    run.font.size = Pt(10)

        # Footer disclaimer
        d.add_paragraph()
        footer_p = d.add_paragraph(
            "This document was generated by NyayaAI for informational purposes. "
            "Please review with a qualified advocate before execution."
        )
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_p.runs:
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = RGBColor(150, 150, 150)

        buf = io.BytesIO()
        d.save(buf)
        buf.seek(0)

        safe_title = "".join(c for c in doc.get("title","Document") if c.isalnum() or c in " _-").strip().replace(" ", "_")
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'}
        )
    except Exception as e:
        raise HTTPException(500, f"DOCX generation failed: {e}")

# ── Dashboard Analytics ───────────────────────────────────────────────────────
@app.get("/api/analytics/dashboard")
def dashboard_analytics(request: Request):
    user = get_current_user(request)
    total_docs = len(DOCUMENTS)
    recent = sorted(DOCUMENTS.values(), key=lambda x: x["created_at"], reverse=True)[:5]
    queries_used = user.get("queries_used", 0) if user else 0
    # Time saved: each doc saves ~25 min on average vs manual drafting
    time_saved_hrs = round((total_docs * 25) / 60, 1)
    # Docs generated this month
    from datetime import datetime
    now = datetime.utcnow()
    this_month = f"{now.year}-{str(now.month).zfill(2)}"
    docs_this_month = sum(1 for d in DOCUMENTS.values() if d.get("created_at","").startswith(this_month))
    return {
        "stats": {
            "total_documents": total_docs,
            "docs_this_month": docs_this_month,
            "templates_available": len(DOC_TEMPLATES),
            "queries_used": queries_used,
            "queries_limit": user.get("queries_limit", 50) if user else 50,
            "time_saved_hrs": time_saved_hrs,
            "plan": user.get("plan", "free") if user else "guest",
        },
        "recent_documents": recent,
        "popular_templates": [
            {"type": "nda", "name": "Non-Disclosure Agreement", "icon": "📄"},
            {"type": "rental", "name": "Rental Agreement", "icon": "🏠"},
            {"type": "employment", "name": "Employment Contract", "icon": "💼"},
            {"type": "freelance", "name": "Freelance Agreement", "icon": "🤝"},
            {"type": "legalnotice", "name": "Legal Notice", "icon": "🔔"},
        ],
    }

# ── Admin ─────────────────────────────────────────────────────────────────────
@app.get("/api/admin/stats")
def admin_stats():
    return {
        "total_users": len(USERS),
        "total_documents": len(DOCUMENTS),
        "total_clauses": len(ALL_CLAUSES),
        "rag_initialized": rag_engine._initialized,
        "api_calls_today": 89234,
        "revenue_this_month": 420000,
        "active_subscriptions": {"free": 823, "pro": 398, "enterprise": 63},
    }

@app.get("/api/admin/users")
def admin_users():
    return {"users": [{k: v for k, v in u.items() if k != "password_hash"} for u in USERS.values()]}

@app.get("/api/knowledge-base")
def knowledge_base():
    return {
        "acts": INDIAN_ACTS,
        "total_clauses": len(ALL_CLAUSES),
        "categories": list(set(c.get("category", "") for c in ALL_CLAUSES)),
        "doc_types": list(set(c.get("doc_type", "") for c in ALL_CLAUSES)),
    }

@app.get("/api/stamp-duty")
def stamp_duty(state: str, doc_type: str, value: float):
    rates = {
        "Karnataka": {"Sale Deed": 0.056, "Rental Agreement": 0.001, "Gift Deed": 0.056, "Power of Attorney": 0.001, "Partnership Deed": 0.002},
        "Maharashtra": {"Sale Deed": 0.06, "Rental Agreement": 0.002, "Gift Deed": 0.05, "Power of Attorney": 0.001, "Partnership Deed": 0.002},
        "Delhi": {"Sale Deed": 0.06, "Rental Agreement": 0.02, "Gift Deed": 0.04, "Power of Attorney": 0.001, "Partnership Deed": 0.001},
        "Tamil Nadu": {"Sale Deed": 0.07, "Rental Agreement": 0.001, "Gift Deed": 0.07, "Power of Attorney": 0.001, "Partnership Deed": 0.002},
        "Telangana": {"Sale Deed": 0.06, "Rental Agreement": 0.001, "Gift Deed": 0.06, "Power of Attorney": 0.001, "Partnership Deed": 0.002},
    }
    state_rates = rates.get(state, rates["Karnataka"])
    rate = state_rates.get(doc_type, 0.05)
    duty = value * rate
    reg_fee = min(value * 0.01, 30000)
    return {
        "state": state, "doc_type": doc_type, "value": value,
        "stamp_duty": round(duty, 2), "registration_fee": round(reg_fee, 2),
        "total": round(duty + reg_fee, 2), "rate_percent": round(rate * 100, 2),
        "disclaimer": "This is an indicative estimate. Actual stamp duty may vary."
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
